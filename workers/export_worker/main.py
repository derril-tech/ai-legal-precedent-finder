"""
Export Worker - Main Entry Point

Generates exports: briefs (DOCX/PDF), citation tables (CSV), JSON bundles.
Handles:
- Document composition
- Citation formatting
- File generation and storage
- Signed URL generation
"""

import asyncio
import logging
from typing import Dict, Any, List, Optional
import json
import csv
import io
from datetime import datetime
import uuid

import nats
from nats.aio.client import Client as NATS
from nats.aio.msg import Msg
from sqlalchemy import create_engine, text
from sqlalchemy.orm import sessionmaker
from fastapi import FastAPI
import boto3
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .brief_generator import BriefGenerator
from .citation_table_generator import CitationTableGenerator

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# FastAPI app for health checks
app = FastAPI(title="Export Worker", version="0.1.0")

# Configuration
DATABASE_URL = "postgresql://postgres:postgres@postgres:5432/legal_precedent_finder"
NATS_URL = "nats://nats:4222"
S3_ENDPOINT = "http://minio:9000"
S3_ACCESS_KEY = "minioadmin"
S3_SECRET_KEY = "minioadmin"

# Initialize components
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

s3_client = boto3.client(
    's3',
    endpoint_url=S3_ENDPOINT,
    aws_access_key_id=S3_ACCESS_KEY,
    aws_secret_access_key=S3_SECRET_KEY
)

brief_generator = BriefGenerator()
citation_table_generator = CitationTableGenerator()

async def process_export_make(msg: Msg):
    """Process export make message from NATS"""
    try:
        data = json.loads(msg.data.decode())
        logger.info(f"Processing export make: {data.get('export_id', 'unknown')}")
        
        export_id = data.get('export_id')
        export_type = data.get('export_type')  # 'brief', 'citations', 'json'
        session_id = data.get('session_id')
        workspace_id = data.get('workspace_id')
        
        # Generate export
        result = await generate_export(export_id, export_type, session_id, workspace_id)
        
        # Publish completion event
        await publish_export_completed(result)
        
        # Acknowledge message
        await msg.ack()
        
    except Exception as e:
        logger.error(f"Error processing export make: {e}")
        await msg.nak()

async def generate_export(
    export_id: str,
    export_type: str,
    session_id: str,
    workspace_id: str
) -> Dict[str, Any]:
    """
    Generate export based on type
    
    Args:
        export_id: Export identifier
        export_type: Type of export ('brief', 'citations', 'json')
        session_id: QA session identifier
        workspace_id: Workspace context
        
    Returns:
        Dictionary with export results
    """
    try:
        with SessionLocal() as session:
            # Get QA session and answer data
            session_query = text("""
                SELECT qs.id, qs.question, a.answer_text, a.reasoning
                FROM qa_sessions qs
                JOIN answers a ON qs.id = a.qa_session_id
                WHERE qs.id = :session_id AND qs.workspace_id = :workspace_id
            """)
            
            session_result = session.execute(session_query, {
                'session_id': session_id,
                'workspace_id': workspace_id
            })
            session_data = session_result.fetchone()
            
            if not session_data:
                raise ValueError(f"QA session {session_id} not found")
            
            # Get citations
            citations_query = text("""
                SELECT ac.citation_text, c.case_name, c.citation, c.court, 
                       c.jurisdiction, ac.relevance_score
                FROM answer_citations ac
                JOIN cases c ON ac.case_id = c.id
                JOIN answers a ON ac.answer_id = a.id
                WHERE a.qa_session_id = :session_id
                ORDER BY ac.relevance_score DESC
            """)
            
            citations_result = session.execute(citations_query, {'session_id': session_id})
            citations = citations_result.fetchall()
            
            # Generate export based on type
            if export_type == 'brief':
                result = await generate_brief(
                    export_id=export_id,
                    question=session_data.question,
                    answer=session_data.answer_text,
                    reasoning=session_data.reasoning,
                    citations=citations
                )
            elif export_type == 'citations':
                result = await generate_citation_table(
                    export_id=export_id,
                    citations=citations
                )
            elif export_type == 'json':
                result = await generate_json_bundle(
                    export_id=export_id,
                    session_data=session_data,
                    citations=citations
                )
            else:
                raise ValueError(f"Unsupported export type: {export_type}")
            
            # Store export metadata
            await store_export_metadata(
                session=session,
                export_id=export_id,
                export_type=export_type,
                session_id=session_id,
                result=result
            )
            
            session.commit()
            
            return result
            
    except Exception as e:
        logger.error(f"Error generating export {export_id}: {e}")
        return {
            'export_id': export_id,
            'status': 'error',
            'error': str(e)
        }

async def generate_brief(
    export_id: str,
    question: str,
    answer: str,
    reasoning: str,
    citations: List
) -> Dict[str, Any]:
    """Generate legal brief document"""
    
    # Create DOCX document
    doc = Document()
    
    # Title
    title = doc.add_heading('Legal Research Brief', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")}')
    doc.add_paragraph(f'Export ID: {export_id}')
    doc.add_paragraph('')
    
    # Question
    doc.add_heading('Question', level=1)
    doc.add_paragraph(question)
    doc.add_paragraph('')
    
    # Answer
    doc.add_heading('Answer', level=1)
    doc.add_paragraph(answer)
    doc.add_paragraph('')
    
    # Reasoning
    if reasoning:
        doc.add_heading('Reasoning', level=1)
        doc.add_paragraph(reasoning)
        doc.add_paragraph('')
    
    # Citations
    if citations:
        doc.add_heading('Supporting Authorities', level=1)
        for citation in citations:
            p = doc.add_paragraph()
            p.add_run(f'{citation.case_name}, ').bold = True
            p.add_run(f'{citation.citation} ({citation.court})')
            p.add_run(f' - Relevance: {citation.relevance_score:.2f}')
    
    # Disclaimer
    doc.add_paragraph('')
    disclaimer = doc.add_paragraph(
        'DISCLAIMER: This brief is generated for research purposes only and does not constitute legal advice. '
        'Please consult with a qualified attorney for specific legal guidance.'
    )
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Upload to S3
    s3_key = f"exports/{export_id}/brief.docx"
    s3_client.put_object(
        Bucket="legal-exports",
        Key=s3_key,
        Body=buffer.getvalue(),
        ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Generate signed URL
    signed_url = s3_client.generate_presigned_url(
        'get_object',
        Params={'Bucket': 'legal-exports', 'Key': s3_key},
        ExpiresIn=3600  # 1 hour
    )
    
    return {
        'export_id': export_id,
        'status': 'success',
        'export_type': 'brief',
        'file_format': 'docx',
        's3_key': s3_key,
        'signed_url': signed_url,
        'citations_count': len(citations)
    }

async def generate_citation_table(
    export_id: str,
    citations: List
) -> Dict[str, Any]:
    """Generate citation table CSV"""
    
    # Create CSV buffer
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    
    # Write header
    writer.writerow([
        'Case Name',
        'Citation',
        'Court',
        'Jurisdiction',
        'Relevance Score',
        'Citation Text'
    ])
    
    # Write data
    for citation in citations:
        writer.writerow([
            citation.case_name,
            citation.citation,
            citation.court,
            citation.jurisdiction,
            citation.relevance_score,
            citation.citation_text
        ])
    
    csv_content = buffer.getvalue()
    
    # Upload to S3
    s3_key = f"exports/{export_id}/citations.csv"
    s3_client.put_object(
        Bucket="legal-exports",
        Key=s3_key,
        Body=csv_content.encode('utf-8'),
        ContentType="text/csv"
    )
    
    # Generate signed URL
    signed_url = s3_client.generate_presigned_url(
        'get_object',
        Params={'Bucket': 'legal-exports', 'Key': s3_key},
        ExpiresIn=3600  # 1 hour
    )
    
    return {
        'export_id': export_id,
        'status': 'success',
        'export_type': 'citations',
        'file_format': 'csv',
        's3_key': s3_key,
        'signed_url': signed_url,
        'citations_count': len(citations)
    }

async def generate_json_bundle(
    export_id: str,
    session_data,
    citations: List
) -> Dict[str, Any]:
    """Generate JSON bundle with all data"""
    
    # Create JSON bundle
    bundle = {
        'export_id': export_id,
        'generated_at': datetime.utcnow().isoformat(),
        'question': session_data.question,
        'answer': session_data.answer_text,
        'reasoning': session_data.reasoning,
        'citations': [
            {
                'case_name': c.case_name,
                'citation': c.citation,
                'court': c.court,
                'jurisdiction': c.jurisdiction,
                'relevance_score': float(c.relevance_score),
                'citation_text': c.citation_text
            }
            for c in citations
        ]
    }
    
    json_content = json.dumps(bundle, indent=2)
    
    # Upload to S3
    s3_key = f"exports/{export_id}/bundle.json"
    s3_client.put_object(
        Bucket="legal-exports",
        Key=s3_key,
        Body=json_content.encode('utf-8'),
        ContentType="application/json"
    )
    
    # Generate signed URL
    signed_url = s3_client.generate_presigned_url(
        'get_object',
        Params={'Bucket': 'legal-exports', 'Key': s3_key},
        ExpiresIn=3600  # 1 hour
    )
    
    return {
        'export_id': export_id,
        'status': 'success',
        'export_type': 'json',
        'file_format': 'json',
        's3_key': s3_key,
        'signed_url': signed_url,
        'citations_count': len(citations)
    }

async def store_export_metadata(
    session,
    export_id: str,
    export_type: str,
    session_id: str,
    result: Dict[str, Any]
):
    """Store export metadata in database"""
    
    # This would store export metadata in a new exports table
    # For now, we'll log it
    logger.info(f"Export {export_id} completed: {result}")

async def publish_export_completed(result: Dict[str, Any]):
    """Publish export completed event"""
    # This will be implemented when NATS client is available
    logger.info(f"Publishing export completed event: {result}")

# FastAPI endpoints
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "service": "export-worker"}

@app.post("/exports/generate")
async def manual_generate_export(
    export_type: str,
    session_id: str,
    workspace_id: str
):
    """Manual export generation endpoint"""
    export_id = str(uuid.uuid4())
    result = await generate_export(export_id, export_type, session_id, workspace_id)
    return result

async def main():
    """Main entry point"""
    
    # Connect to NATS
    nc = NATS()
    await nc.connect(NATS_URL)
    
    # Subscribe to export make events
    await nc.subscribe("export.make", cb=process_export_make)
    
    logger.info("Export worker started and listening for events")
    
    # Keep the worker running
    try:
        while True:
            await asyncio.sleep(1)
    except KeyboardInterrupt:
        await nc.close()

if __name__ == "__main__":
    asyncio.run(main())
